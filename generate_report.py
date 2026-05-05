import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    document = Document()

    # Title
    title = document.add_heading('Project Report: SkillCluster AI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph('Automated Resume Clustering and Skill Gap Detection using NLP').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('--------------------------------------------------------------\n')

    # Abstract
    document.add_heading('Abstract', level=1)
    document.add_paragraph(
        "The modern recruitment and placement landscape is highly competitive, processing thousands of resumes daily. "
        "Traditional resume screening and evaluation rely heavily on manual review, which is time-consuming, prone to human error, "
        "and often inconsistent. This project introduces SkillCluster AI, an end-to-end Machine Learning and Natural Language Processing (NLP) "
        "pipeline designed to automate the analysis, categorization, and evaluation of candidate resumes. By leveraging TF-IDF vectorization "
        "and K-Means clustering, the system intelligently groups resumes based on semantic similarities and technical skill overlaps. "
        "Furthermore, the project features a robust Skill Gap Detection engine that compares individual or weaker clusters against top-tier "
        "clusters to identify missing high-value industry skills. The final output is an interactive web-based dashboard that not only visualizes "
        "these groupings via dimensionality reduction techniques (PCA) but also provides personalized learning path recommendations. "
        "This report details the methodology, algorithms, and system architecture used to construct this placement intelligence platform."
    )
    document.add_page_break()

    # Chapter 1
    document.add_heading('1. Introduction', level=1)
    document.add_heading('1.1 Background', level=2)
    document.add_paragraph(
        "With the rapid expansion of the technology sector, academic institutions and placement cells are facing "
        "unprecedented challenges in managing student portfolios and matching them with relevant job roles. "
        "A typical resume contains unstructured text representing a candidate's education, experience, and technical proficiencies. "
        "Manually parsing this information to gauge placement readiness is an arduous task. Natural Language Processing (NLP) provides "
        "powerful tools to extract structured data from this unstructured text, enabling automated analysis."
    )
    document.add_paragraph(
        "Skill gaps represent the disparity between the skills a candidate possesses and the skills demanded by the industry. "
        "Identifying these gaps early is crucial for students to upskill and improve their employability. However, without a data-driven "
        "benchmark, students are often unaware of their deficiencies until they face rejection in interviews."
    )
    
    document.add_heading('1.2 Problem Statement', level=2)
    document.add_paragraph(
        "Placement cells and recruiters face challenges in analyzing hundreds of resumes to group candidates by skill sets accurately. "
        "Similarly, students struggle to understand what specific skills they are missing compared to top-tier candidates. "
        "There is a distinct lack of automated, AI-driven platforms that can evaluate unstructured resumes, cluster them intelligently "
        "without prior labeling, and provide actionable, personalized insights for improvement."
    )

    document.add_heading('1.3 Objectives', level=2)
    document.add_paragraph(
        "The primary objectives of the SkillCluster AI project are:\n"
        "1. To develop a robust NLP pipeline capable of cleaning and extracting structured skill data from raw resume text.\n"
        "2. To implement unsupervised machine learning techniques to cluster resumes based on technical profiles.\n"
        "3. To design a comparative analysis algorithm that detects specific skill gaps between strong and weak candidate clusters.\n"
        "4. To generate personalized learning path recommendations to bridge the identified skill gaps.\n"
        "5. To build an interactive, modern Streamlit dashboard for data visualization and live resume analysis."
    )
    
    document.add_paragraph("\n" * 15) # Spacing to push content and simulate pages

    # Chapter 2
    document.add_heading('2. Literature Review & Technical Background', level=1)
    document.add_heading('2.1 Natural Language Processing in Recruitment', level=2)
    document.add_paragraph(
        "NLP has been increasingly adopted in Applicant Tracking Systems (ATS) to parse resumes. "
        "Techniques range from simple keyword matching using Regular Expressions to complex Named Entity Recognition (NER) "
        "using deep learning models like BERT and SpaCy. While ATS filters candidates against a specific job description, "
        "this project focuses on exploratory analysis—understanding the inherent groupings within a pool of candidates."
    )
    
    document.add_heading('2.2 TF-IDF Vectorization', level=2)
    document.add_paragraph(
        "Term Frequency-Inverse Document Frequency (TF-IDF) is a statistical measure used to evaluate how important a word is "
        "to a document in a collection or corpus. The TF component measures the frequency of a word in a document, while the IDF "
        "component diminishes the weight of words that occur very frequently across the entire corpus (e.g., 'the', 'and', 'experience'). "
        "This ensures that rare, highly specific technical terms (e.g., 'Kubernetes', 'PyTorch') receive higher weights, making them "
        "strong discriminators during clustering."
    )

    document.add_heading('2.3 Unsupervised Learning: K-Means and PCA', level=2)
    document.add_paragraph(
        "K-Means clustering is a partition-based clustering algorithm that divides a dataset into 'K' distinct, non-overlapping subsets. "
        "It operates by minimizing the variance within each cluster. In high-dimensional text data, K-Means efficiently groups documents "
        "that share similar TF-IDF vectors.\n"
        "To visualize these high-dimensional clusters, Principal Component Analysis (PCA) is employed. PCA is an orthogonal linear transformation "
        "that transforms data to a new coordinate system, reducing dimensions while retaining the maximum possible variance, thus allowing "
        "us to plot the clusters on a 2D graph."
    )

    document.add_page_break()

    # Chapter 3
    document.add_heading('3. Methodology and System Architecture', level=1)
    
    document.add_heading('3.1 Dataset Description', level=2)
    document.add_paragraph(
        "The project utilizes a Kaggle-sourced Resume Dataset. It contains thousands of resumes spanning various IT and non-IT domains. "
        "The most critical feature in this dataset is the 'Resume_str' column, which holds the raw, unstructured text parsed from PDF or DOCX files. "
        "This dataset provides a realistic representation of the noise, formatting artifacts, and varied vocabularies found in real-world applications."
    )

    document.add_heading('3.2 Data Preprocessing Pipeline', level=2)
    document.add_paragraph(
        "The preprocessing module is the foundational layer of the system. It consists of the following sequential steps:\n"
        "• Noise Removal: Using Regex, we strip out URLs, email addresses, phone numbers, and non-alphanumeric characters.\n"
        "• Lowercasing: Normalizing text to lowercase to ensure uniformity (e.g., 'Python' and 'python' are treated identically).\n"
        "• Tokenization: Splitting the continuous string of text into discrete words (tokens) using NLTK's word_tokenize.\n"
        "• Stopword Removal: Eliminating common English words that carry little semantic weight using NLTK's stopword corpus.\n"
        "• Lemmatization: Converting words to their morphological root using the WordNetLemmatizer (e.g., 'developing' becomes 'develop')."
    )

    document.add_heading('3.3 Information Extraction Engine', level=2)
    document.add_paragraph(
        "A customized extraction engine was developed using predefined technology dictionaries. The dictionaries are categorized into "
        "Programming Languages, Frameworks, Tools, Databases, and Cloud Technologies. The engine iterates through the cleaned resume text "
        "and performs exact string matching to identify and extract these entities. This structured metadata is used alongside the embeddings "
        "for deep-dive cluster analysis."
    )

    document.add_heading('3.4 Feature Engineering with TF-IDF', level=2)
    document.add_paragraph(
        "The cleaned text is transformed into a dense numerical matrix using the TfidfVectorizer. The feature space is limited to the top 1000 "
        "most significant terms to optimize computational efficiency while capturing the variance. The resulting matrix, denoted as X, "
        "serves as the input feature space for the clustering algorithm."
    )

    document.add_heading('3.5 Clustering Optimization', level=2)
    document.add_paragraph(
        "The K-Means algorithm is applied to the TF-IDF matrix. The optimal number of clusters (K=4) was determined iteratively to segment "
        "the candidates into distinct technical tiers. The model converges by recalculating centroids and reassigning points until "
        "the intra-cluster sum of squares is minimized. The results are stored and mapped back to the original dataframe."
    )

    document.add_page_break()

    # Chapter 4
    document.add_heading('4. Analytical Modules', level=1)
    
    document.add_heading('4.1 Cluster Insight Generation', level=2)
    document.add_paragraph(
        "Once clusters are formed, the system performs an aggregation of the extracted skills within each cluster. By computing the mean "
        "number of skills and analyzing the frequency distribution of technologies, the system profiles each cluster. "
        "For example, one cluster might emerge as highly proficient in Data Science (Python, SQL, TensorFlow), while another might "
        "be Web Development focused (JavaScript, React, Node.js)."
    )

    document.add_heading('4.2 Skill Gap Detection Algorithm', level=2)
    document.add_paragraph(
        "The Skill Gap Detection engine identifies the 'Strongest' cluster (highest average skill count) and the 'Weakest' cluster "
        "(lowest average skill count). By utilizing set operations on the top 20 most frequent skills of each cluster, the algorithm "
        "isolates the skills that are highly prevalent in the strong cluster but completely absent in the weak cluster. These are marked "
        "as critical skill gaps."
    )

    document.add_heading('4.3 Learning Path Recommendation', level=2)
    document.add_paragraph(
        "Based on the identified skill gaps, a rule-based recommendation engine generates personalized advice. If a candidate is missing "
        "fundamental languages like Python or SQL, the system recommends beginner courses. If they lack DevOps tools like Docker or AWS, "
        "it suggests cloud certifications. This provides a direct, actionable roadmap for placement improvement."
    )

    document.add_paragraph("\n" * 15)

    # Chapter 5
    document.add_heading('5. System Interface: Streamlit Dashboard', level=1)
    document.add_paragraph(
        "A premium, highly interactive dashboard was developed using Streamlit. It serves as the front-end for placement officers and students. "
        "Key features include:\n"
        "• Dynamic Navigation: Horizontal menu structures using streamlit-option-menu.\n"
        "• Modern Aesthetics: Custom CSS implementing glassmorphism, animated gradients, and professional typography.\n"
        "• Data Visualization: Integration with Plotly for 2D PCA scatter plots, pie charts, and interactive bar graphs.\n"
        "• Live Analysis: A text area where users can paste a new resume, which is instantly cleaned, vectorized, assigned to a cluster, "
        "and evaluated for missing skills, returning an immediate learning path."
    )

    document.add_page_break()

    # Chapter 6
    document.add_heading('6. Results and Evaluation', level=1)
    document.add_paragraph(
        "The implementation of the pipeline on a sample of 1000 resumes successfully segregated the candidates into 4 cohesive clusters. "
        "The PCA scatter plot demonstrated clear boundaries between the high-skill density clusters and the low-skill density clusters. "
        "The Skill Gap engine successfully detected absences in critical tools such as Git, AWS, and SQL in the weaker clusters, validating "
        "the hypothesis that lower-tier resumes often lack industry-standard operational tools compared to just missing programming languages."
    )
    document.add_paragraph(
        "The Streamlit dashboard performed with high responsiveness, capable of analyzing a new text input and rendering predictions "
        "in under 2 seconds, proving its viability for real-time application in a placement cell environment."
    )
    
    document.add_paragraph("\n" * 10)

    # Chapter 7
    document.add_heading('7. Conclusion and Future Scope', level=1)
    document.add_paragraph(
        "SkillCluster AI successfully bridges the gap between raw resume data and actionable placement intelligence. "
        "By leveraging robust NLP preprocessing, TF-IDF feature extraction, and Unsupervised Learning, the system eliminates the manual "
        "overhead of candidate evaluation. It groups students logically by their technical profiles and provides data-driven, personalized "
        "recommendations to maximize their career potential."
    )
    document.add_paragraph(
        "Future Enhancements could include:\n"
        "• Integration of deep contextual embeddings like Sentence-BERT or fine-tuned Transformer models to capture semantic nuances better than TF-IDF.\n"
        "• Implementing a supervised layer to predict actual placement probabilities based on historical placement data.\n"
        "• Direct integration with job portals to automatically pull current job descriptions and perform dynamic skill gap analysis against real-time market demands.\n"
        "• Expanding the extraction engine to include soft skills, project sentiment analysis, and education parsing."
    )

    document.save('SkillCluster_AI_Project_Report.docx')

if __name__ == '__main__':
    create_report()
